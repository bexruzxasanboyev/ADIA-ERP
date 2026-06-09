# -*- coding: utf-8 -*-
"""Generate ADIA ERP TZ + muddat-hisobi .docx (Uzbek, Latin). One-off."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

OUT = os.path.dirname(os.path.abspath(__file__))
DARK = RGBColor(0x1F, 0x2A, 0x44)
ACCENT = RGBColor(0x2D, 0x6C, 0xDF)
GREEN = RGBColor(0x1B, 0x7F, 0x3B)
AMBER = RGBColor(0xB4, 0x6A, 0x00)
RED = RGBColor(0xB0, 0x2A, 0x37)
GREY = RGBColor(0x55, 0x5F, 0x6D)

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def style_doc(doc):
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)

# ---------------------------------------------------------------------------
# Module data: (no, name, status, desc, reqs[], dep, accept, days)
# status: 'Tayyor' | 'Qisman' | 'Yangi'
MODULES = [
 (1, "Bo'limlar bog'liqligi (ta'minot zanjiri)", "Tayyor",
  "Barcha bo'g'inlar (xom-ashyo ombori → ishlab chiqarish sexlari → sex skladlari → markaziy sklad → do'konlar) bir-biriga bog'lanadi. Ostatka min'dan tushganda tizim avtomatik to'ldirish (replenishment) tsiklini ishga tushiradi.",
  ["Har (product, location) uchun min/max chegaralari; dinamik qayta hisob",
   "Replenishment so'rovining holat-mashinasi (10 holat) va audit izi",
   "Har 5 daqiqada avtomatik skan (cron) — kam qolgan pozitsiyalarni topadi",
   "Bitta (product, location) uchun bir vaqtda faqat bitta ochiq so'rov (debounce)"],
  "—", "Min'dan past pozitsiya avtomatik so'rov yaratadi; so'rov zanjir bo'ylab to'g'ri o'tadi; har o'tish audit-logga yoziladi.", 1),

 (2, "Ishlab chiqarish ombori — zagatovka logikasi", "Tayyor",
  "Sex o'z skladiga qarab: tayyor yarim-fabrikat (zagatovka) bor bo'lsa undan foydalanadi; bo'lmasa retseptdan kelib chiqib ingredientlarni mahsulot omboridan so'raydi. Bir necha tayyor mahsulot uchun kerakli xom-ashyo (masalan krem, shakar) umumiy holda jamlanib, bitta so'rov sifatida mahsulot omboriga boradi.",
  ["Retsept bosqichlari: asos (hamir/zagatovka), bezak (krem+ukrasheniye), yig'ish",
   "Ishlab chiqarish buyrug'ida 'tayyordan yoki 0dan?' tekshiruvi",
   "Sex skladidagi mavjud zagatovka birinchi tekshiriladi",
   "Bir necha mahsulot uchun ingredient so'rovlari jamlanadi (agregatsiya)"],
  "M1, M13", "Zagatovka tayyor bo'lsa hamir qayta so'ralmaydi; yetishmagan material bo'yicha jamlangan so'rov yaratiladi; barchasi bitta atomar tranzaksiyada.", 2),

 (3, "Mahsulot bo'limi bog'lanishi", "Tayyor",
  "Mahsulot (ma'lumotnoma) bo'limi markaziy sklad mantig'i bilan bir xil tarzda zanjirga ulanadi.",
  ["Mahsulotlar ro'yxati: xom-ashyo / yarim tayyor / tayyor bo'yicha tablar",
   "Kategoriya, sex va birlik bo'yicha filtr; tannarx (computed_cost) ko'rsatish",
   "Retsept (BOM) ko'rish va tahrirlash (sikl tekshiruvi bilan)"],
  "M4", "Mahsulotlar to'g'ri turkumlanadi; narx retseptdan hisoblanadi; markaziy sklad oqimlariga ulangan.", 1),

 (4, "Mahsulot bo'limi — Poster mahsulotlariga bog'lanishi", "Tayyor",
  "Mahsulotlar Poster POS bilan moslanadi: Ингредиенты → xom-ashyo, Полуфабрикаты → yarim tayyor/tayyor, тех.карта → kategoriya+rasm+sex manbai. Ingredientlar mahsulot omboriga tortiladi.",
  ["Poster'dan idempotent sinxronizatsiya (ikki marta yurish dublikat yaratmaydi)",
   "«Г/П» prefiksli polufabrikatlar — tayyor mahsulot; qolganlari — yarim tayyor",
   "Retsept (brutto/netto, birlik normalizatsiyasi: g→kg, ml→l) import",
   "Товары (Coca Cola va sh.k.) mahsulot sifatida olinmaydi"],
  "M3", "Sinxron natijasi: xom-ashyo/yarim/tayyor to'g'ri ajraladi; товары yo'q; retseptlar to'g'ri import bo'ladi.", 2),

 (5, "Sexlar o'rtasida bog'lanish", "Tayyor",
  "Sexlar bir-biriga material/yarim-mahsulot uzatadi; katta mahsulotlar bir necha sex tomonidan qo'shma ishlab chiqariladi.",
  ["Sexlararo oqim grafigi (location_flows): production_output / bom_input / forward / reverse",
   "Ko'p bosqichli retsept: zagatovka bir sexda, bezak boshqa sexda",
   "Oqimlarni boshqarish (admin UI) va vizualizatsiya"],
  "M1, M2", "Sexlararo uzatish oqimlari aniq; ko'p sexli mahsulot to'g'ri yig'iladi.", 1),

 (6, "Qaymoq krem ishlab chiqarish bo'limi", "Qisman",
  "Retseptlarning aksariyatida qaymoq krem bor — u ko'p ishlatilgani uchun alohida bo'lim/sex sifatida ajratiladi. Bu sex boshqa sexlardan so'rovnoma oladi, krem tayyorlaydi va yetkazadi. Sexning o'z zagatovka skladi bo'ladi.",
  ["Alohida 'Qaymoq/Krem' ishlab chiqarish sexi (bo'g'in) sifatida sozlash",
   "Boshqa sexlardan krem so'rovnomasini qabul qilish",
   "Krem ishlab chiqarib, so'ragan sexga yetkazish (oqim)",
   "Krem bo'limining o'z zagatovka skladi va ostatkasi"],
  "M2, M5", "Boshqa sex krem so'raydi → krem bo'limi so'rovni oladi → tayyorlaydi → so'ragan sexga uzatadi; ostatka to'g'ri yangilanadi.", 7),

 (7, "Ishlab chiqarish KPI", "Tayyor",
  "Har mahsulot uchun tannarx + komunal (umumiy xarajat) + oylik summalar asosida so'mda KPI; oy oxirida umumiy oylik (ish haqi) qo'shilib foyda hisoblanadi.",
  ["Tannarx (material) — retsept rollupidan; komunal — birlik uchun",
   "Ish haqi ulushi — ishlab chiqarish bo'limi oyliklari / ishlab chiqarilgan birlik",
   "To'liq tannarx = material + komunal + ish haqi; foyda = tushum − to'liq tannarx",
   "Oylik hisobot va KPI-maqsad (kpi_target) bo'yicha taqqoslash"],
  "M3, M4", "Oylik KPI hisoboti har mahsulot bo'yicha to'liq tannarx va foydani ko'rsatadi.", 2),

 (8, "Do'kon KPI", "Qisman",
  "Sotuvchi/do'kon bo'yicha sotuv soniga plan va KPI qo'yiladi; shunga qarab sotuvlarni oshirish kuzatiladi.",
  ["Do'kon/sotuvchi bo'yicha oylik sotuv PLANI (maqsad) jadvali — YANGI",
   "Plan vs haqiqiy bajarilish (% maqsad) ko'rsatkichi — YANGI",
   "Sotuvchi darajasidagi kuzatuv (sotuvni sotuvchiga bog'lash) — YANGI",
   "O'sish (week/month) tahlili va reyting"],
  "M10", "Har do'kon/sotuvchi uchun plan qo'yiladi; bajarilish foizi va o'sish ko'rinadi; ortda qolganlar ajralib turadi.", 8),

 (9, "Kassa tafovuti / fors-major ogohlantirishlar", "Qisman",
  "Do'konda Poster va ERP qoldig'i mos kelmaganda (qoldiq 0 bo'lsa-da sotuv urilsa yoki ortiqcha sotuv) — real-vaqt rejimida aniqlanib, bot 'noto'g'ri chek' deb ogohlantiradi.",
  ["Ortiqcha sotuv / manfiy ostatka aniqlash (sinxron tsiklida — mavjud)",
   "Per-do'kon jamlangan Telegram ogohlantirish (mavjud)",
   "Real-vaqtga yaqin (1 daqiqalik webhook) tezligini sozlash va kuchaytirish — QISMAN",
   "Ogohlantirish UI/hisoboti (noto'g'ri cheklar ro'yxati)"],
  "M10", "Ortiqcha/xato sotuv 1 daqiqa ichida aniqlanadi; mas'ul shaxsga jamlangan ogohlantirish boradi.", 4),

 (10, "Inventarizatsiya — ERP↔Poster ostatka sinxron", "Tayyor",
  "Poster'dagi umumiy qoldiq ERP bilan avtomatik sinxronlanadi (avto shakllantirish).",
  ["Poster storage.getStorageLeftovers → ERP stock (har 15 daqiqada cron)",
   "Farq bo'lsa qabul/chiqim harakati yaratiladi; ostatka manfiy bo'lmaydi (0 ga clamp)",
   "Manfiy qoldiq bo'yicha per-location jamlangan ogohlantirish"],
  "M4", "Poster qoldig'i ERP'ga aks etadi; farqlar harakat sifatida yoziladi; manfiy qoldiq ogohlantiriladi.", 1),

 (11, "Inventarizatsiya konverteri/kalkulyatori (bo'lak↔butun)", "Yangi",
  "Tort/Napoleon kabi mahsulotlar BUTUN ishlab chiqariladi, lekin BO'LAK qilib sotiladi. Kun oxirida masalan tortning yarmi sotilib yarmi qolsa — inventarizatsiya buni to'g'ri hisoblashi kerak (bo'lak sotuvlarni butun birlikka konvert qilish).",
  ["Mahsulot uchun 'butun → necha bo'lak' koeffitsiyenti (recipe_yield asosida)",
   "Bo'lak (kusok) sotuvlarni butun (целый) birlikka konvert qilish kalkulyatori",
   "Kun oxiri inventarizatsiya: qisman sotilgan butun mahsulotni hisoblash",
   "Konverter natijasini ostatka bilan moslab inventarizatsiya yozuvi"],
  "M10", "Napoleon butun ishlab chiqariladi; 3 bo'lak sotilib 7 qolsa — tizim qoldiqni to'g'ri (butun+bo'lak) hisoblaydi.", 7),

 (12, "Yangi mahsulot ishlab chiqarish", "Tayyor",
  "Poster'da yo'q (polufabrikat/mahsulot mavjud bo'lmagan) yangi mahsulot uchun ishlab chiqarish: so'rov yuboriladi, retsept asosida mahsulot ombori material beradi va mahsulot yangi mahsulot sifatida sotuvga olib boriladi.",
  ["Poster'da yo'q yangi mahsulot yaratish (nom, tur, birlik, retsept)",
   "Yangi mahsulotga ishlab chiqarish buyrug'i berish",
   "Retsept asosida material so'rovi va atomar ishlab chiqarish (consume+produce)"],
  "M2, M13, M14", "Yangi mahsulot yaratiladi → so'rov → material beriladi → ishlab chiqariladi → sotuvga tayyor.", 1),

 (13, "Ishlab chiqarish so'rovi", "Tayyor",
  "Ishlab chiqarish buyrug'i/so'rovi yaratish va tasdiqlash oqimi; AI-dialog orqali manba tanlash.",
  ["Ishlab chiqarish buyrug'i: new → in_progress → done (yoki cancelled)",
   "AI-dialog: Q1 'tayyordan yoki 0dan?', Q2 'krem tayyorlash yoki ombordan?'",
   "Web + Telegram kanallarida bir xil dialog; 6 soatda muddati o'tsa PM'ga eskalatsiya"],
  "M2", "So'rov yaratiladi → dialog savollariga javob → kerakli hujjatlar bitta tranzaksiyada yaratiladi.", 1),

 (14, "Mahsulot ombori so'rovi", "Tayyor",
  "Mahsulot/xom-ashyo omboriga so'rov; boshliq + skladchi ikki bosqichli tasdiq.",
  ["So'rov yaratish (ta'minot menejeri yoki admin)",
   "Ikki bosqichli tasdiq: boshliq + skladchi — ikkalasi ham tasdiqlashi shart",
   "Qabul qilish: stock harakati + brak (defekt) qty/sababi",
   "RBAC bo'yicha ko'rinish (har rol o'z so'rovini ko'radi)"],
  "M1", "So'rov ikki tasdiqdan keyin kuchga kiradi; qabulda brak hisobga olinadi; ostatka yangilanadi.", 1),

 (15, "Kassir boti", "Qisman",
  "Telegram kassir boti: Poster seyf (pul hisoblari) va cheklardan foydalanib, kun oxirida hisobotlarni topshiradi.",
  ["Erkin matndan smena ma'lumotini o'qish (rasxod/qoldiq/karta) — mavjud",
   "Pul-asosli nakladnoy (kun-oxiri hisobot) yaratish — mavjud",
   "Poster seyf/finance hisoblaridan balansni o'qish (finance.*) — QISMAN/YANGI",
   "Cheklardan kun-oxiri jamlanma va solishtirish"],
  "M9, M10", "Kassir Telegram orqali kun-oxiri hisobotini topshiradi; Poster seyf balansi bilan solishtiriladi.", 5),
]

PHASES = [
 ("Bosqich 1 — Ishlab chiqarish yadrosi", "Asosan tayyor; yakuniy moslash va sinov (UAT).", [1,2,5,13,12]),
 ("Bosqich 2 — Mahsulot, ombor va Poster", "Tayyor; verifikatsiya va integratsiya sinovi.", [3,4,14,10]),
 ("Bosqich 3 — KPI bloki", "Ishlab chiqarish KPI tayyor; Do'kon KPI yangi quriladi.", [7,8]),
 ("Bosqich 4 — Yangi/qisman modullar", "Asosiy yangi ish: krem bo'limi, inventarizatsiya konverteri, kassa tafovuti.", [6,11,9]),
 ("Bosqich 5 — Kassir bot va yakun", "Kassir bot Poster integratsiyasi + yakuniy.", [15]),
]
EXTRA = [("Integratsiya testi + xato tuzatish buferi", 5),
         ("Deploy + UAT (egasi bilan qabul) + hujjatlashtirish", 3)]

def status_color(s):
    return {'Tayyor':GREEN,'Qisman':AMBER,'Yangi':RED}.get(s,GREY)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = DARK
    return h

# ===========================================================================
# DOC 1 — TZ
# ===========================================================================
def build_tz():
    doc = Document(); style_doc(doc)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ADIA ERP"); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = ACCENT
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Texnik Topshiriq (TZ)"); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = DARK
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Non/tort/qandolat ishlab chiqarish va ta'minot zanjiri uchun o'zini-o'zi to'g'rilaydigan ERP tizimi")
    r.italic = True; r.font.color.rgb = GREY
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run("Sana: " + datetime.date.today().strftime("%d.%m.%Y")).font.color.rgb = GREY
    doc.add_paragraph()

    add_heading(doc, "1. Kirish", 1)
    doc.add_paragraph(
        "ADIA ERP — xom-ashyo omboridan boshlab sexlar, sex skladlari, markaziy "
        "sklad va do'konlargacha bo'lgan butun ta'minot zanjirini boshqaruvchi tizim. "
        "Tizimning asosiy g'oyasi: har mahsulotga minimal/maksimal chegara belgilanadi; "
        "ostatka minimaldan tushganda tizim avtomatik to'ldirish tsiklini ishga tushiradi. "
        "Savdo va ombor ma'lumotlari Poster POS platformasidan sinxronlanadi. "
        "Eng yuqori qatlamda — AI yordamchi (Vertex AI Gemini).")

    add_heading(doc, "2. Maqsad va qamrov", 1)
    doc.add_paragraph(
        "Ushbu TZ quyidagi 15 ta funksional modulni qamrab oladi. Har modul uchun maqsad, "
        "funksional talablar, bog'liqliklar, qabul mezonlari va joriy holat keltirilgan. "
        "Joriy holat uch xil belgilanadi: «Tayyor» (ishlab chiqarilgan), «Qisman» (asosi bor, "
        "to'ldirish kerak), «Yangi» (noldan quriladi).")

    add_heading(doc, "3. Texnik stek", 1)
    for line in ["Frontend: React + Vite + TypeScript; shadcn/ui + Tailwind (dark premium); Recharts",
                 "Backend: Node.js + Express; raw SQL qatlami; PostgreSQL",
                 "Auth: JWT + RBAC (rolga asoslangan ruxsat)",
                 "Fon ishlar: node-cron (replenishment skan, sinxronizatsiya)",
                 "Bot: Telegram (Grammy)",
                 "AI: Vertex AI Gemini (DB ustida function calling)",
                 "Integratsiya: Poster POS API (savdo, cheklar, ombor qoldig'i)",
                 "Deploy: Hetzner VPS · PM2 · Nginx"]:
        p = doc.add_paragraph(style='List Bullet'); p.add_run(line)

    add_heading(doc, "4. Joriy holat (umumiy)", 1)
    built = sum(1 for m in MODULES if m[2]=='Tayyor'); part = sum(1 for m in MODULES if m[2]=='Qisman'); new = sum(1 for m in MODULES if m[2]=='Yangi')
    doc.add_paragraph(
        f"Jami 15 ta moduldan: «Tayyor» — {built} ta, «Qisman» — {part} ta, «Yangi» — {new} ta. "
        "Ya'ni tizimning asosiy yadrosi (ta'minot zanjiri, ishlab chiqarish, Poster integratsiyasi, "
        "narx/KPI) allaqachon ishlab chiqilgan; qolgan ish asosan Do'kon KPI, qaymoq krem bo'limi, "
        "inventarizatsiya konverteri va kassa tafovuti modullariga to'g'ri keladi.")

    add_heading(doc, "5. Modullar", 1)
    for no,name,status,desc,reqs,dep,acc,days in MODULES:
        h = doc.add_heading(level=2);
        run = h.add_run(f"5.{no}. {name}  "); run.font.color.rgb = DARK
        badge = h.add_run(f"[{status}]"); badge.font.color.rgb = status_color(status); badge.bold = True
        doc.add_paragraph(desc)
        p = doc.add_paragraph(); p.add_run("Funksional talablar:").bold = True
        for rq in reqs:
            b = doc.add_paragraph(style='List Bullet'); b.add_run(rq)
        p = doc.add_paragraph(); p.add_run("Bog'liqliklar: ").bold = True; p.add_run(dep if dep else "—")
        p = doc.add_paragraph(); p.add_run("Qabul mezoni: ").bold = True; p.add_run(acc)

    add_heading(doc, "6. Nofunksional talablar", 1)
    for line in ["Atomarlik: har stock harakati atomar tranzaksiya (yo hammasi, yo hech narsa).",
                 "Ostatka hech qachon manfiy bo'lmaydi (DB CHECK + ilova tekshiruvi).",
                 "RBAC qat'iy: har rol faqat o'z bo'g'inini ko'radi; PM/Admin butun zanjirni.",
                 "Har o'zgarish audit-logga yoziladi (kim, qachon, nima).",
                 "Poster sinxronizatsiyasi idempotent (qayta yurish dublikat yaratmaydi).",
                 "Narxlash app-owned: xom-ashyo narxi qo'lda, hosila narxlar retseptdan hisoblanadi.",
                 "Ikki bosqichli tasdiq: ta'minot so'rovi boshliq+skladchi tasdig'idan keyin kuchga kiradi."]:
        p = doc.add_paragraph(style='List Bullet'); p.add_run(line)

    add_heading(doc, "7. Bosqichlar", 1)
    doc.add_paragraph("Yetkazib berish quyidagi bosqichlarga bo'linadi (batafsil muddat alohida hujjatda):")
    for title, note, ids in PHASES:
        p = doc.add_paragraph(style='List Bullet'); p.add_run(title + " — ").bold = True; p.add_run(note)

    doc.save(os.path.join(OUT, "ADIA_ERP_TZ.docx"))
    print("TZ saved")

# ===========================================================================
# DOC 2 — Muddat hisobi
# ===========================================================================
def build_estimate():
    doc = Document(); style_doc(doc)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ADIA ERP — Muddat Hisobi"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = ACCENT
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Texnik topshiriq bo'yicha taxminiy ish muddati"); r.italic = True; r.font.color.rgb = GREY
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run("Sana: " + datetime.date.today().strftime("%d.%m.%Y")).font.color.rgb = GREY
    doc.add_paragraph()

    add_heading(doc, "Asoslar (farazlar)", 1)
    for line in ["1 ta tajribali full-stack dasturchi; ish kuni = 8 soat; ish haftasi = 5 kun.",
                 "Muddat — sof ishlab chiqish + sinov; bayram/ta'til hisobga olinmagan.",
                 "Tizimning katta qismi allaqachon ishlab chiqilgan — bu umumiy muddatni qisqartiradi.",
                 "«Tayyor» modullar uchun muddat — yakuniy moslash, integratsiya va qabul sinovi (UAT).",
                 "Muddatlar taxminiy; aniq talablar oydinlashganda ±20% o'zgarishi mumkin."]:
        p = doc.add_paragraph(style='List Bullet'); p.add_run(line)

    add_heading(doc, "Modullar bo'yicha muddat", 1)
    tbl = doc.add_table(rows=1, cols=4); tbl.style = 'Light Grid Accent 1'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i,htext in enumerate(["#  Modul","Holat","Qolgan ish","Ish kuni"]):
        hdr[i].paragraphs[0].add_run(htext).bold = True
    for no,name,status,desc,reqs,dep,acc,days in MODULES:
        c = tbl.add_row().cells
        c[0].paragraphs[0].add_run(f"{no}. {name}")
        rr = c[1].paragraphs[0].add_run(status); rr.font.color.rgb = status_color(status); rr.bold = True
        gap = "Yakuniy moslash + UAT" if status=='Tayyor' else ("To'ldirish + integratsiya" if status=='Qisman' else "Noldan qurish")
        c[2].paragraphs[0].add_run(gap)
        c[3].paragraphs[0].add_run(str(days)); c[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    mod_total = sum(m[7] for m in MODULES)

    add_heading(doc, "Bosqichlar bo'yicha jamlanma", 1)
    tbl2 = doc.add_table(rows=1, cols=3); tbl2.style = 'Light Grid Accent 1'
    h2 = tbl2.rows[0].cells
    for i,htext in enumerate(["Bosqich","Modullar","Ish kuni"]):
        h2[i].paragraphs[0].add_run(htext).bold = True
    grand = 0
    for title, note, ids in PHASES:
        days = sum(next(m[7] for m in MODULES if m[0]==i) for i in ids); grand += days
        c = tbl2.add_row().cells
        c[0].paragraphs[0].add_run(title)
        c[1].paragraphs[0].add_run(", ".join(str(i) for i in ids))
        c[2].paragraphs[0].add_run(str(days)); c[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, days in EXTRA:
        grand += days
        c = tbl2.add_row().cells
        c[0].paragraphs[0].add_run(label); c[1].paragraphs[0].add_run("—")
        c[2].paragraphs[0].add_run(str(days)); c[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # total row
    c = tbl2.add_row().cells
    c[0].paragraphs[0].add_run("JAMI").bold = True; c[1].paragraphs[0].add_run("")
    tr = c[2].paragraphs[0].add_run(str(grand)); tr.bold = True; c[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell in c: set_cell_bg(cell, "E8EEFB")

    weeks = round(grand/5,1)
    add_heading(doc, "Umumiy muddat", 1)
    p = doc.add_paragraph()
    p.add_run(f"Jami taxminan {grand} ish kuni ").bold = True
    p.add_run(f"(≈ {weeks} ish haftasi, ≈ {round(grand/21,1)} oy — 1 dasturchi uchun).")
    doc.add_paragraph(
        "Eslatma: tizimning ~70% allaqachon ishlab chiqilgan. Yangi/qisman modullar "
        "(Do'kon KPI, Qaymoq krem bo'limi, Inventarizatsiya konverteri, Kassa tafovuti, Kassir bot) "
        "asosiy qolgan ishni tashkil etadi. Bir vaqtda 2 dasturchi ishlasa, umumiy muddat "
        "sezilarli qisqaradi.")

    doc.save(os.path.join(OUT, "ADIA_ERP_Muddat_Hisobi.docx"))
    print(f"Estimate saved; module-sum={mod_total}, grand={grand} days, {weeks} weeks")

build_tz()
build_estimate()
print("DONE ->", OUT)
